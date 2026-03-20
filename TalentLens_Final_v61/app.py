# TalentLens — Post-Resume Talent Intelligence
# Techkriti '26 x Eightfold AI | Impact Area 01: Signal Extraction & Verification
#
# Datasets:
#   Kaggle Resume Dataset  -> skill vocabulary (snehaanbhawal/resume-dataset)
#   LinkedIn Job Postings  -> job search & requirements (arshkon/linkedin-job-postings)
#   Lightcast Open Skills  -> taxonomy, lang->skill map (data/skills_taxonomy.json)

import streamlit as st
import html
from sentence_transformers import SentenceTransformer, util
import torch
import ast
import traceback
from io import StringIO, BytesIO
import contextlib
import time
import re
import json
import requests
import pandas as pd
from collections import defaultdict, Counter
from pathlib import Path
import base64

# Resume file parsers
try:
    import fitz as pymupdf          # PyMuPDF — PDF text extraction
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document as DocxDocument  # python-docx — DOCX extraction
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from PIL import Image            # Pillow — image handling for JPG/PNG
    PIL_OK = True
except ImportError:
    PIL_OK = False

st.set_page_config(
    page_title="TalentLens — Signal Intelligence",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Styles

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;1,9..40,400&family=DM+Mono:wght@400;500&display=swap');

html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; color: #1a1a2e; }
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 1.5rem 2.5rem 4rem 2.5rem !important; max-width: 1200px !important; }

.hero {
    background: linear-gradient(135deg, #0f172a 0%, #1e1b4b 60%, #0f172a 100%);
    border-radius: 20px; padding: 2.5rem 3rem; margin-bottom: 2rem;
    position: relative; overflow: hidden;
}
.hero::before {
    content:''; position:absolute; top:-60px; right:-60px;
    width:300px; height:300px; border-radius:50%;
    background:radial-gradient(circle,rgba(99,102,241,.25) 0%,transparent 70%);
}
.hero-eyebrow { font-size:.72rem; font-weight:600; letter-spacing:.14em; text-transform:uppercase; color:#818cf8; margin-bottom:.75rem; }
.hero-title   { font-size:2.2rem; font-weight:700; color:#f8fafc; line-height:1.2; margin:0 0 .5rem 0; }
.hero-title span { color:#818cf8; }
.hero-sub     { font-size:.95rem; color:#94a3b8; max-width:640px; margin:0; }
.hero-chips   { display:flex; gap:.6rem; flex-wrap:wrap; margin-top:1.25rem; }
.hero-chip    { background:rgba(99,102,241,.15); border:1px solid rgba(99,102,241,.3); color:#a5b4fc; font-size:.75rem; font-weight:600; padding:.25rem .75rem; border-radius:20px; }

.sec-label { font-size:.7rem; font-weight:700; letter-spacing:.14em; text-transform:uppercase; color:#94a3b8; display:flex; align-items:center; gap:.5rem; margin-bottom:.5rem; }
.sec-label::after { content:''; flex:1; height:1px; background:#e5e7eb; }
.sec-title { font-size:1.15rem; font-weight:700; color:#111827; margin-bottom:.2rem; }
.sec-desc  { font-size:.86rem; color:#6b7280; margin-bottom:1.1rem; }

.card { background:#fff; border:1px solid #e5e7eb; border-radius:14px; padding:1.5rem; margin-bottom:1.25rem; box-shadow:0 1px 3px rgba(0,0,0,.04); }
.job-card { background:#f8faff; border:1px solid #c7d2fe; border-radius:12px; padding:1.1rem 1.25rem; cursor:pointer; transition:border-color .2s; margin-bottom:.5rem; }
.job-card:hover { border-color:#4f46e5; }
.job-card-title { font-weight:700; font-size:.92rem; color:#111827; }
.job-card-meta  { font-size:.78rem; color:#6b7280; margin-top:.2rem; }

.evidence-item { display:flex; gap:.75rem; padding:.9rem 1rem; border-radius:10px; margin-bottom:.5rem; border:1px solid #e5e7eb; background:#fafafa; align-items:flex-start; }
.evidence-claim { font-size:.88rem; font-weight:600; color:#111827; }
.evidence-because { font-size:.82rem; color:#6b7280; margin-top:.15rem; line-height:1.6; }
.ev-src { display:inline-flex; align-items:center; gap:.25rem; font-size:.72rem; font-weight:600; padding:.1rem .45rem; border-radius:20px; margin:.2rem .2rem 0 0; }
.ev-src.github  { background:#0d1117; color:#58a6ff; border:1px solid #30363d; }
.ev-src.resume  { background:#f0fdf4; color:#16a34a; border:1px solid #bbf7d0; }
.ev-src.task    { background:#faf5ff; color:#7c3aed; border:1px solid #ddd6fe; }
.ev-src.semantic{ background:#fff7ed; color:#c2410c; border:1px solid #fed7aa; }
.ev-src.dataset { background:#eff6ff; color:#1d4ed8; border:1px solid #bfdbfe; }

.lang-bar-bg { background:#21262d; border-radius:4px; height:8px; overflow:hidden; display:flex; }
.lang-seg { height:100%; }
.lang-legend { display:flex; flex-wrap:wrap; gap:.5rem; margin-top:.4rem; }
.lang-dot { display:inline-flex; align-items:center; gap:.3rem; font-size:.75rem; color:#e6edf3; }

.pill-row { display:flex; flex-wrap:wrap; gap:.4rem; margin-top:.4rem; }
.pill { display:inline-flex; align-items:center; gap:.3rem; padding:.25rem .65rem; border-radius:20px; font-size:.8rem; font-weight:500; }
.pill.matched    { background:#f0fdf4; color:#15803d; border:1px solid #bbf7d0; }
.pill.missing    { background:#fff7ed; color:#c2410c; border:1px solid #fed7aa; }
.pill.verified   { background:#eff6ff; color:#1d4ed8; border:1px solid #bfdbfe; }
.pill.unverified { background:#f9fafb; color:#6b7280; border:1px solid #e5e7eb; }
.pill.github-p   { background:#0d1117; color:#58a6ff; border:1px solid #30363d; }
.pill.dataset-p  { background:#eff6ff; color:#1d4ed8; border:1px solid #bfdbfe; }

.ring-wrap { position:relative; width:100px; height:100px; }
.ring-wrap svg { transform:rotate(-90deg); }
.ring-inner { position:absolute; inset:0; display:flex; flex-direction:column; align-items:center; justify-content:center; }
.ring-num   { font-size:1.7rem; font-weight:700; color:#111827; line-height:1; }
.ring-denom { font-size:.7rem; color:#9ca3af; font-weight:500; }

.stat-card { background:#f9fafb; border:1px solid #e5e7eb; border-radius:12px; padding:1rem 1.25rem; }
.stat-val  { font-size:1.5rem; font-weight:700; color:#111827; line-height:1; margin-bottom:.2rem; }
.stat-lab  { font-size:.72rem; color:#9ca3af; font-weight:500; text-transform:uppercase; letter-spacing:.06em; }

.feedback-mono { background:#f8faff; border:1px solid #dde6ff; border-left:4px solid #4f46e5; border-radius:0 10px 10px 0; padding:1rem 1.25rem; font-family:'DM Mono',monospace; font-size:.8rem; line-height:1.7; color:#374151; white-space:pre-wrap; }
.mentor-block  { background:linear-gradient(135deg,#faf5ff,#f5f3ff); border:1px solid #ddd6fe; border-left:4px solid #7c3aed; border-radius:0 10px 10px 0; padding:1rem 1.25rem; font-size:.88rem; line-height:1.8; color:#374151; }
.explain-card  { background:linear-gradient(135deg,#fefce8,#fef9c3); border:1px solid #fde68a; border-radius:14px; padding:1.5rem; margin-bottom:1rem; }
.bias-card     { background:linear-gradient(135deg,#f0fdf4,#dcfce7); border:1px solid #86efac; border-radius:14px; padding:1.5rem; margin-bottom:1rem; }

.final-card  { background:linear-gradient(135deg,#0f172a,#1e1b4b); border:1px solid #312e81; border-radius:16px; padding:2rem; }
.final-score { font-size:3.5rem; font-weight:800; line-height:1; }
.badge { display:inline-flex; align-items:center; padding:.3rem .8rem; border-radius:20px; font-size:.8rem; font-weight:600; }
.badge.high   { background:#dcfce7; color:#15803d; }
.badge.medium { background:#fef9c3; color:#ca8a04; }
.badge.low    { background:#fee2e2; color:#dc2626; }
.badge.behavior { background:#ede9fe; color:#6d28d9; }

.task-card { background:#f8faff; border:1px solid #c7d2fe; border-radius:14px; padding:1.5rem; margin-bottom:1rem; }
.divider   { height:1px; background:linear-gradient(90deg,transparent,#e5e7eb 20%,#e5e7eb 80%,transparent); margin:2rem 0; }

.dataset-badge { display:inline-flex; align-items:center; gap:.3rem; background:#eff6ff; color:#1d4ed8; border:1px solid #bfdbfe; font-size:.68rem; font-weight:700; padding:.15rem .5rem; border-radius:20px; vertical-align:middle; margin-left:.3rem; }
.llm-badge     { display:inline-flex; align-items:center; gap:.3rem; background:linear-gradient(135deg,#4f46e5,#7c3aed); color:white; font-size:.68rem; font-weight:700; padding:.15rem .5rem; border-radius:20px; vertical-align:middle; margin-left:.3rem; }

.stTextArea textarea { font-family:'DM Mono',monospace!important; font-size:.84rem!important; border-radius:10px!important; border-color:#e5e7eb!important; background:#fafafa!important; }
.stTextArea textarea:focus { border-color:#4f46e5!important; box-shadow:0 0 0 3px rgba(79,70,229,.08)!important; }
.stTextInput input { border-radius:10px!important; border-color:#e5e7eb!important; }
.stButton > button { background:#4f46e5!important; color:#fff!important; border:none!important; border-radius:10px!important; font-family:'DM Sans',sans-serif!important; font-size:.88rem!important; font-weight:600!important; padding:.55rem 1.5rem!important; box-shadow:0 2px 8px rgba(79,70,229,.2)!important; }
.stButton > button:hover { background:#4338ca!important; }
label { font-weight:500!important; font-size:.86rem!important; color:#374151!important; }
</style>
""", unsafe_allow_html=True)

# Config

DATA_DIR    = Path(__file__).parent / "data"
CLAUDE_MODEL = "claude-haiku-4-5-20251001"
LANG_COLORS  = {
    "Python":"#3572A5","JavaScript":"#f1e05a","TypeScript":"#2b7489",
    "Java":"#b07219","C++":"#f34b7d","C":"#555555","Go":"#00ADD8",
    "Rust":"#dea584","Ruby":"#701516","Shell":"#89e051","Kotlin":"#F18E33",
    "Swift":"#ffac45","R":"#198CE7","Jupyter Notebook":"#DA5B0B",
    "HTML":"#e34c26","CSS":"#563d7c","Scala":"#c22d40","PHP":"#4F5D95",
    "Dart":"#00B4AB","Solidity":"#AA6746","C#":"#178600","MATLAB":"#e16737",
    "default":"#8b5cf6",
}

# Dataset loading

@st.cache_data(show_spinner=False)
def load_skills_taxonomy() -> dict:
    """Load the bundled Lightcast skills taxonomy."""
    try:
        with open(DATA_DIR / "skills_taxonomy.json") as f:
            return json.load(f)
    except Exception:
        return {"categories": {}, "language_to_skills": {}, "skill_aliases": {}, "job_role_skills": {}}

@st.cache_data(show_spinner=False)
def load_resume_skill_vocabulary() -> set:
    """Build skill vocabulary from Kaggle Resume Dataset + Lightcast taxonomy.

    If Resume.csv is present it boosts the vocab with real resume tokens.
    Falls back to taxonomy only when the CSV is missing.
    """
    taxonomy = load_skills_taxonomy()
    
    # Build base vocabulary from taxonomy (always available)
    vocab = set()
    for skills in taxonomy.get("categories", {}).values():
        vocab.update(s.lower() for s in skills)
    for alias in taxonomy.get("skill_aliases", {}).keys():
        vocab.add(alias.lower())

    # Augment with Resume.csv if available
    resume_path = DATA_DIR / "Resume.csv"
    if resume_path.exists():
        try:
            df = pd.read_csv(resume_path)
            text_col = next((c for c in df.columns if "resume" in c.lower()), None)
            if text_col:
                all_text = " ".join(df[text_col].dropna().astype(str).tolist()).lower()
                # Extract CamelCase tokens and known tech patterns
                tokens = re.findall(r'\b[A-Z][a-z]+(?:[A-Z][a-z]+)+\b', " ".join(df[text_col].dropna()))
                token_counts = Counter(t.lower() for t in tokens)
                # Add tokens that appear 2+ times (dataset-validated skills)
                vocab.update(t for t, c in token_counts.items() if c >= 2)
        except Exception:
            pass  # Graceful degradation to taxonomy-only

    return vocab

@st.cache_data(show_spinner=False)
def load_job_postings() -> pd.DataFrame:
    """Load the LinkedIn Job Postings CSV if available."""
    postings_path = DATA_DIR / "postings.csv"
    if postings_path.exists():
        try:
            df = pd.read_csv(postings_path)
            # Normalise column names to handle both original and sample schemas
            df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
            return df
        except Exception:
            return pd.DataFrame()
    return pd.DataFrame()

@st.cache_resource(show_spinner="Loading embedding model…")
def load_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

taxonomy      = load_skills_taxonomy()
skill_vocab   = load_resume_skill_vocabulary()
jobs_df       = load_job_postings()
model         = load_model()

# Skill extraction

def dataset_extract_skills(text: str) -> list:
    """
    Extract skills using vocabulary derived from Kaggle Resume Dataset +
    Lightcast taxonomy. Zero hardcoded skill keywords — all terms come from
    the datasets loaded at startup.
    """
    found = []
    text_lower = text.lower()

    # Taxonomy pass
    for category, skills in taxonomy.get("categories", {}).items():
        for skill in skills:
            pattern = rf"\b{re.escape(skill.lower())}\b"
            if re.search(pattern, text_lower):
                found.append(skill)

    # Alias resolution pass
    aliases = taxonomy.get("skill_aliases", {})
    for alias, canonical in aliases.items():
        if re.search(rf"\b{re.escape(alias.lower())}\b", text_lower):
            if canonical not in found:
                found.append(canonical)

    # Resume vocab pass
    for token in skill_vocab:
        if len(token) >= 3 and token not in [s.lower() for s in found]:
            if re.search(rf"\b{re.escape(token)}\b", text_lower):
                # Capitalise correctly using taxonomy lookup
                canonical = next(
                    (s for skills in taxonomy.get("categories", {}).values()
                     for s in skills if s.lower() == token),
                    token.title()
                )
                found.append(canonical)

    return list(dict.fromkeys(found))

# Job posting search

def search_job_postings(query: str, max_results: int = 6) -> list:
    """
    Search the LinkedIn Job Postings dataset for relevant roles.
    Returns list of job dicts with title, company, skills, description.
    Source: https://www.kaggle.com/datasets/arshkon/linkedin-job-postings
    """
    if jobs_df.empty:
        return []

    query_lower = query.lower()
    results = []

    title_col   = next((c for c in jobs_df.columns if "title"   in c), None)
    skills_col  = next((c for c in jobs_df.columns if "skill"   in c), None)
    company_col = next((c for c in jobs_df.columns if "company" in c), None)
    desc_col    = next((c for c in jobs_df.columns if "desc"    in c), None)
    loc_col     = next((c for c in jobs_df.columns if "location" in c), None)

    if not title_col:
        return []

    for _, row in jobs_df.iterrows():
        title      = str(row.get(title_col, "")).strip()
        skills_raw = str(row.get(skills_col, "") or "")
        company    = str(row.get(company_col, "") or "")
        desc       = str(row.get(desc_col, "") or "")
        location   = str(row.get(loc_col, "") or "")

        searchable = f"{title} {company} {skills_raw} {desc}".lower()
        if query_lower in searchable or any(w in searchable for w in query_lower.split()):
            skills_list = [s.strip() for s in skills_raw.split(",") if s.strip()]
            results.append({
                "title":    title,
                "company":  company,
                "skills":   skills_list,
                "location": location,
                "desc_preview": (desc[:120] + "…") if len(desc) > 120 else desc,
            })
        if len(results) >= max_results:
            break

    return results

def get_skills_for_role(role_title: str) -> list:
    """
    Get required skills for a role from the LinkedIn Job Postings dataset
    or fall back to taxonomy job_role_skills.
    """
    postings = search_job_postings(role_title, max_results=1)
    if postings and postings[0]["skills"]:
        return postings[0]["skills"]

    role_skills = taxonomy.get("job_role_skills", {})
    for role, skills in role_skills.items():
        if role_title.lower() in role.lower() or role.lower() in role_title.lower():
            return skills

    return []

# Claude API helper

def _call_claude(api_key: str, prompt: str, max_tokens: int = 600,
                 system: str = "") -> str:
    body = {"model": CLAUDE_MODEL, "max_tokens": max_tokens,
            "messages": [{"role": "user", "content": prompt}]}
    if system:
        body["system"] = system
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                 "content-type": "application/json"},
        json=body, timeout=25,
    )
    resp.raise_for_status()
    return resp.json()["content"][0]["text"].strip()

def _strip_fences(t: str) -> str:
    return re.sub(r"^```(?:json)?\s*|\s*```$", "", t, flags=re.MULTILINE).strip()

# LLM skill extraction

def llm_extract_skills(text: str, api_key: str) -> list:
    """Extract skills from text via Claude NER, grounded in Lightcast taxonomy."""
    # Get a sample of taxonomy terms to anchor the extraction
    sample_skills = []
    for skills in list(taxonomy.get("categories", {}).values())[:5]:
        sample_skills.extend(skills[:4])

    prompt = (
        f"Extract ALL technical and professional skills from the text below.\n"
        f"Use standardised skill names matching industry taxonomy "
        f"(e.g. examples: {', '.join(sample_skills[:15])}).\n"
        f"Include: programming languages, frameworks, tools, platforms, methodologies, soft skills.\n"
        f"Return ONLY a JSON array of strings. No markdown, no preamble.\n"
        f'Text:\n"""\n{text[:3000]}\n"""'
    )
    try:
        raw = _call_claude(api_key, prompt, max_tokens=512)
        skills = [s.strip() for s in json.loads(_strip_fences(raw)) if isinstance(s, str)]
        # Resolve aliases using taxonomy
        aliases = taxonomy.get("skill_aliases", {})
        resolved = []
        for s in skills:
            resolved.append(aliases.get(s, s))
        return list(dict.fromkeys(resolved))
    except Exception:
        return dataset_extract_skills(text)

# GitHub profile parser

def parse_github_profile(username: str, gh_token: str = "") -> dict:
    headers = {"Accept": "application/vnd.github+json"}
    if gh_token:
        headers["Authorization"] = f"Bearer {gh_token}"
    base = "https://api.github.com"
    signals = {
        "username": username, "languages": {}, "top_repos": [],
        "total_stars": 0, "total_repos": 0, "push_events": 0,
        "inferred_skills": [], "raw_languages": {}, "error": None, "bio": "",
    }

    try:
        u = requests.get(f"{base}/users/{username}", headers=headers, timeout=10)
        if u.status_code == 404:
            signals["error"] = f"GitHub user '{username}' not found."
            return signals
        u.raise_for_status()
        ud = u.json()
        signals["total_repos"] = ud.get("public_repos", 0)
        signals["followers"]   = ud.get("followers", 0)
        signals["bio"]         = ud.get("bio") or ""
    except Exception as e:
        signals["error"] = str(e)
        return signals

    try:
        r = requests.get(f"{base}/users/{username}/repos",
                         headers=headers,
                         params={"per_page": 100, "sort": "updated"},
                         timeout=15)
        r.raise_for_status()
        repos = r.json()
    except Exception as e:
        signals["error"] = str(e)
        return signals

    lang_bytes: dict = defaultdict(int)
    for repo in repos:
        if repo.get("fork"):
            continue
        stars = repo.get("stargazers_count", 0)
        signals["total_stars"] += stars
        lang = repo.get("language")
        if lang:
            lang_bytes[lang] += repo.get("size", 1)
        signals["top_repos"].append({
            "name":        repo["name"],
            "description": repo.get("description") or "",
            "stars":       stars,
            "language":    lang or "—",
            "topics":      repo.get("topics", []),
            "url":         repo.get("html_url", ""),
            "updated":     (repo.get("pushed_at") or "")[:10],
        })

    signals["top_repos"] = sorted(signals["top_repos"], key=lambda x: x["stars"], reverse=True)[:8]
    total_bytes = sum(lang_bytes.values()) or 1
    signals["raw_languages"] = dict(lang_bytes)
    signals["languages"] = {
        lang: round(b / total_bytes * 100, 1)
        for lang, b in sorted(lang_bytes.items(), key=lambda x: -x[1])
    }

    try:
        ev = requests.get(f"{base}/users/{username}/events/public",
                          headers=headers, params={"per_page": 100}, timeout=10)
        if ev.status_code == 200:
            signals["push_events"] = sum(1 for e in ev.json() if e.get("type") == "PushEvent")
    except Exception:
        pass

    # Infer skills from GitHub languages + repo topics
    lang_skill_map = taxonomy.get("language_to_skills", {})
    inferred = []
    topic_signals = []
    for repo in signals["top_repos"]:
        topic_signals.extend(repo.get("topics", []))

    for lang, pct in signals["languages"].items():
        if pct >= 5 and lang in lang_skill_map:
            for skill in lang_skill_map[lang][:3]:
                inferred.append({
                    "skill":    skill,
                    "evidence": f"{pct:.0f}% of codebase is {lang}",
                    "source":   "github_language",
                    "strength": "strong" if pct >= 20 else "moderate",
                })

    for topic in set(topic_signals):
        topic_clean = topic.replace("-", " ")
        canonical   = next(
            (s for skills in taxonomy.get("categories", {}).values()
             for s in skills if s.lower() == topic_clean.lower()),
            topic_clean.title()
        )
        inferred.append({
            "skill":    canonical,
            "evidence": f"repository tagged '{topic}'",
            "source":   "github_topic",
            "strength": "moderate",
        })

    for repo in signals["top_repos"][:3]:
        if repo["stars"] >= 5:
            inferred.append({
                "skill":    "Open Source Contribution",
                "evidence": f"'{repo['name']}' — {repo['stars']} ⭐ community-validated",
                "source":   "github_stars",
                "strength": "strong",
            })
            break

    signals["inferred_skills"] = inferred
    return signals

# Semantic skill matching

def match_skills(job_skills: list, candidate_skills: list,
                 threshold: float = 0.55) -> dict:
    if not job_skills or not candidate_skills:
        return {"score": 0, "matched_skills": [], "missing_skills": list(job_skills)}
    j_emb = model.encode(job_skills,       convert_to_tensor=True)
    c_emb = model.encode(candidate_skills, convert_to_tensor=True)
    sim   = util.cos_sim(j_emb, c_emb)
    matched, missing = [], []
    for i, js in enumerate(job_skills):
        best_i = int(torch.argmax(sim[i]).item())
        best_v = float(sim[i][best_i].item())
        if best_v >= threshold:
            matched.append({"job_skill": js,
                            "candidate_skill": candidate_skills[best_i],
                            "similarity": round(best_v, 4)})
        else:
            missing.append(js)
    return {
        "score": round(len(matched) / len(job_skills) * 100),
        "matched_skills": matched,
        "missing_skills": missing,
    }

# Evidence chain builder

def build_evidence_chain(matched_skills, missing_skills,
                          github_signals, llm_skills, task_result) -> list:
    chain         = []
    gh_inferred   = {s["skill"].lower(): s for s in github_signals.get("inferred_skills", [])}
    llm_lower     = {s.lower() for s in llm_skills}
    dataset_skills= {s.lower() for s in dataset_extract_skills(" ".join(llm_skills))}

    for m in matched_skills:
        skill    = m["job_skill"]
        skill_lo = skill.lower()
        sources, because = [], []

        for lang, pct in github_signals.get("languages", {}).items():
            # Check against taxonomy language_to_skills
            lang_skills_lower = [s.lower() for s in taxonomy.get("language_to_skills", {}).get(lang, [])]
            if skill_lo in lang.lower() or lang.lower() in skill_lo or skill_lo in lang_skills_lower:
                sources.append("github")
                because.append(f"{pct:.0f}% of GitHub codebase is {lang}")

        if skill_lo in gh_inferred:
            sources.append("github")
            because.append(gh_inferred[skill_lo]["evidence"])

        for repo in github_signals.get("top_repos", []):
            for topic in repo.get("topics", []):
                if skill_lo in topic.lower() or topic.lower() in skill_lo:
                    sources.append("github")
                    because.append(f"repo '{repo['name']}' tagged '{topic}'")

        matched_word = m.get("candidate_skill", "")
        if matched_word.lower() in llm_lower or skill_lo in llm_lower:
            sources.append("resume")
            because.append(f"extracted from resume/LinkedIn: '{matched_word}'")

        if skill_lo in dataset_skills:
            sources.append("dataset")
            because.append(f"validated against Kaggle Resume Dataset skill vocabulary")

        sim_pct = int(m["similarity"] * 100)
        if not because:
            sources.append("semantic")
            because.append(f"semantic similarity {sim_pct}% to declared '{matched_word}'")

        if task_result and task_result.get("task_score", 0) >= 60:
            if any(t.lower() in skill_lo or skill_lo in t.lower()
                   for t in task_result.get("tags", [])):
                sources.append("task")
                because.append(f"scored {task_result['task_score']}/100 on live coding task")

        chain.append({
            "skill":      skill,
            "status":     "verified",
            "sources":    list(dict.fromkeys(sources)) or ["semantic"],
            "because":    because,
            "confidence": "HIGH" if len(set(sources)) >= 2 else "MEDIUM",
        })

    for s in missing_skills:
        chain.append({
            "skill": s, "status": "missing", "sources": [],
            "because": ["No evidence across GitHub, resume/LinkedIn, or live assessment"],
            "confidence": "LOW",
        })

    return chain

def llm_generate_explanation(chain, candidate_name, job_title, api_key) -> str:
    summary = [
        f"- {e['skill']}: {'VERIFIED' if e['status']=='verified' else 'MISSING'} "
        f"({'; '.join(e['because'][:2])})"
        for e in chain[:10]
    ]
    prompt = (
        f"You are an AI talent analyst. Write a 3-sentence hiring intelligence report "
        f"for '{candidate_name or 'the candidate'}' applying for '{job_title or 'this role'}'.\n\n"
        f"Evidence:\n" + "\n".join(summary) + "\n\n"
        f"1. Overall assessment. 2. Key verified strengths with source. "
        f"3. Gaps and recommendation. Under 100 words. Name specific skills."
    )
    try:
        return _call_claude(api_key, prompt, max_tokens=250)
    except Exception:
        v = [e["skill"] for e in chain if e["status"] == "verified"]
        ms = [e["skill"] for e in chain if e["status"] == "missing"]
        return (f"Candidate has verified evidence for: {', '.join(v[:3])}. "
                f"Gaps: {', '.join(ms[:3]) or 'none'}. Recommend technical interview.")

# Bias audit

def run_bias_audit(job_skills, candidate_skills, demographic_fields) -> dict:
    baseline       = match_skills(job_skills, candidate_skills)
    baseline_score = baseline["score"]
    results = []
    for field, value in demographic_fields.items():
        filtered      = [s for s in candidate_skills
                         if value.lower() not in s.lower() and field.lower() not in s.lower()]
        score_without = match_skills(job_skills, filtered)["score"]
        delta         = abs(score_without - baseline_score)
        results.append({"field": field, "value": value,
                        "score_with": baseline_score, "score_without": score_without,
                        "delta": delta, "pass": delta == 0})
    return {"baseline_score": baseline_score, "checks": results,
            "all_pass": all(r["pass"] for r in results)}

# Task generation

def llm_generate_task(matched_skills: list, job_title: str, api_key: str) -> dict:
    """Generate a coding task based on the candidate's matched skills and role context."""
    skills = [m["job_skill"] for m in matched_skills[:5]]

    role_context = ""
    if job_title:
        role_skills = get_skills_for_role(job_title)
        if role_skills:
            role_context = f"This is for a {job_title} role. Key role skills: {role_skills[:5]}."

    prompt = (
        f"Generate ONE coding task for a candidate with skills: {skills}.\n"
        f"{role_context}\n"
        f"Requirements: single Python function, completable in < 10 min, "
        f"tests practical skill relevant to these technologies, NOT string reversal.\n"
        f"Also generate 3 test cases as a JSON array of [input, expected_output] pairs.\n"
        f"Return ONLY JSON (no markdown):\n"
        f'{{"title":"...","description":"2 clear sentences",'
        f'"placeholder":"def solution(...):\\n    pass",'
        f'"difficulty":"Easy|Medium|Hard","tags":["tag"],'
        f'"evaluation_hint":"solution(x) → y",'
        f'"test_cases":[[input1,output1],[input2,output2],[input3,output3]]}}'
    )
    try:
        raw  = _call_claude(api_key, prompt, max_tokens=600)
        task = json.loads(_strip_fences(raw))
        for k in ("title","description","placeholder","difficulty","tags"):
            if k not in task:
                raise ValueError(k)
        return task
    except Exception:
        return _default_task()

def _default_task() -> dict:
    """Fallback task used when no API key is set."""
    return {
        "title":       "Count Word Frequencies",
        "description": "Write a Python function that takes a list of strings and returns a dictionary with each unique word as a key and its frequency count as the value. The function should be case-insensitive.",
        "placeholder": "def word_frequencies(words: list) -> dict:\n    # your code here\n    pass",
        "difficulty":  "Easy",
        "tags":        ["Python", "Data Structures", "Data Analysis"],
        "evaluation_hint": "word_frequencies(['Apple','apple','Banana']) → {'apple':2,'banana':1}",
        "test_cases":  [
            [["hello","world","hello"], {"hello":2,"world":1}],
            [["Python","python","PYTHON"], {"python":3}],
            [[], {}],
        ],
    }

# Resume file parsers

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyMuPDF (fitz). Handles multi-page resumes."""
    if not PDF_OK:
        return ""
    try:
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        return "\n".join(pages)
    except Exception as e:
        return f"[PDF parse error: {e}]"

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx. Reads all paragraphs and tables."""
    if not DOCX_OK:
        return ""
    try:
        doc = DocxDocument(BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX parse error: {e}]"

def extract_text_from_image(file_bytes: bytes, api_key: str = "") -> str:
    """Use Claude vision to OCR a resume image. Requires API key."""
    if not api_key:
        return ""
    try:
        b64 = base64.standard_b64encode(file_bytes).decode("utf-8")
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            json={
                "model": CLAUDE_MODEL,
                "max_tokens": 1500,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64",
                         "media_type": "image/jpeg", "data": b64}},
                        {"type": "text", "text":
                         "This is a resume image. Extract ALL text content exactly as written. "
                         "Preserve structure — name, contact, skills, experience, education. "
                         "Return plain text only."}
                    ]
                }]
            },
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["content"][0]["text"].strip()
    except Exception as e:
        return f"[Image OCR error: {e}]"

def parse_resume_file(uploaded_file, api_key: str = "") -> tuple[str, str]:
    """Parse uploaded resume file and return (text, parser_label)."""
    if uploaded_file is None:
        return "", ""

    fname = uploaded_file.name.lower()
    raw   = uploaded_file.read()

    if fname.endswith(".pdf"):
        text = extract_text_from_pdf(raw)
        label = "PDF (PyMuPDF)"
    elif fname.endswith(".docx"):
        text = extract_text_from_docx(raw)
        label = "DOCX (python-docx)"
    elif fname.endswith((".jpg", ".jpeg", ".png")):
        text = extract_text_from_image(raw, api_key)
        label = "Image (Claude Vision)"
    else:
        text  = raw.decode("utf-8", errors="ignore")
        label = "Plain text"

    return text.strip(), label

# Multi-question assessment generator

def llm_generate_assessment(matched_skills: list, missing_skills: list,
                              job_title: str, api_key: str) -> dict:
    """
    Generates a full multi-question live assessment grounded in:
      - Candidate's matched + missing skills (from semantic matcher)
      - Job role context from LinkedIn Job Postings dataset
      - Lightcast taxonomy for skill categories

    Returns:
      {
        "coding_task":  { title, description, placeholder, difficulty, tags,
                          evaluation_hint, test_cases },
        "mcq":          [ { question, options[4], answer, skill, explanation } × 4 ],
        "short_answer": [ { question, sample_answer, skill, rubric } × 2 ],
        "total_marks":  100
      }
    """
    skill_names   = [m["job_skill"] for m in matched_skills[:5]]
    missing_names = missing_skills[:3]

    # Role context from dataset
    role_context = ""
    if job_title:
        role_skills = get_skills_for_role(job_title)
        if role_skills:
            role_context = f"Role: {job_title}. Key required skills from dataset: {role_skills[:6]}."

    # Taxonomy skill categories for the matched skills
    skill_categories = []
    for skill in skill_names:
        for cat, skills in taxonomy.get("categories", {}).items():
            if skill in skills:
                skill_categories.append(cat)
                break
    category_hint = f"Skill categories: {list(set(skill_categories))[:3]}." if skill_categories else ""

    prompt = f"""You are a senior technical interviewer. Generate a complete multi-format assessment.

Candidate verified skills: {skill_names}
Skill gaps (missing): {missing_names}
{role_context}
{category_hint}

Generate a JSON object with exactly this structure (no markdown):
{{
  "coding_task": {{
    "title": "...",
    "description": "2 clear sentences explaining the task",
    "placeholder": "def solution(...):\\n    # your code here\\n    pass",
    "difficulty": "Medium",
    "tags": ["skill1", "skill2"],
    "evaluation_hint": "solution(x) → y",
    "test_cases": [[input1, expected1], [input2, expected2], [input3, expected3]]
  }},
  "mcq": [
    {{
      "question": "...",
      "options": ["A) ...", "B) ...", "C) ...", "D) ..."],
      "answer": "A",
      "skill": "skill_being_tested",
      "explanation": "1 sentence why the answer is correct"
    }}
  ],
  "short_answer": [
    {{
      "question": "...",
      "sample_answer": "2-3 sentence model answer",
      "skill": "skill_being_tested",
      "rubric": "what a good answer must mention"
    }}
  ]
}}

Rules:
- coding_task: single Python function, practical, NOT string reversal
- mcq: exactly 4 questions, test the 4 most important matched skills  
- short_answer: exactly 2 questions, one testing a verified skill, one on a gap skill
- All questions must be genuinely distinct and non-trivial
"""
    try:
        raw        = _call_claude(api_key, prompt, max_tokens=1800)
        assessment = json.loads(_strip_fences(raw))
        # Validate structure
        assert "coding_task" in assessment
        assert "mcq" in assessment and len(assessment["mcq"]) >= 2
        assert "short_answer" in assessment and len(assessment["short_answer"]) >= 1
        assessment["total_marks"] = 100
        return assessment
    except Exception:
        # Fall back to coding-only assessment
        return {
            "coding_task":  _default_task(),
            "mcq":          [],
            "short_answer": [],
            "total_marks":  100,
        }

def score_short_answer(question: str, sample_answer: str,
                        rubric: str, candidate_answer: str, api_key: str) -> dict:
    """
    Use Claude to evaluate a candidate's short-answer response against the rubric.
    Returns { score (0-10), feedback, strengths, improvements }
    """
    if not candidate_answer.strip():
        return {"score": 0, "feedback": "No answer provided.",
                "strengths": [], "improvements": ["Please provide an answer."]}
    prompt = f"""Score this short answer out of 10.

Question: {question}
Model answer: {sample_answer}
Rubric (must mention): {rubric}
Candidate's answer: {candidate_answer}

Return ONLY JSON (no markdown):
{{"score": 0-10, "feedback": "2 sentences", "strengths": ["..."], "improvements": ["..."]}}"""
    try:
        raw = _call_claude(api_key, prompt, max_tokens=300)
        return json.loads(_strip_fences(raw))
    except Exception:
        # Fallback: simple keyword match
        keywords = rubric.lower().split()
        hits = sum(1 for k in keywords if k in candidate_answer.lower())
        score = min(10, round(hits / max(len(keywords), 1) * 10))
        return {"score": score, "feedback": "Auto-scored via keyword match.",
                "strengths": [], "improvements": []}

# Code evaluator

def evaluate_code(user_code: str, test_cases: list) -> dict:
    lines, score = [], 0
    try:
        tree = ast.parse(user_code)
    except SyntaxError as exc:
        return {"task_score": 0, "feedback": f"✗ Syntax error: {exc.msg} (line {exc.lineno})"}

    lines.append("✓ Code parsed — no syntax errors.\n")

    ns = {}
    try:
        exec(compile(tree, "<code>", "exec"), ns)
        callables = {k: v for k, v in ns.items() if callable(v) and not k.startswith("_")}
        if not callables:
            raise ValueError("No function found.")
        for hint in ("solution","word_freq","most_freq","count","solve","main"):
            for name, fn in callables.items():
                if hint in name.lower():
                    target_fn = fn; break
            else:
                continue
            break
        else:
            target_fn = next(iter(callables.values()))
    except Exception as exc:
        return {"task_score": 0, "feedback": f"✗ Could not extract function: {exc}"}

    passed, failures = 0, []
    for tc in test_cases:
        if not isinstance(tc, (list, tuple)) or len(tc) < 2:
            continue
        inp, expected = tc[0], tc[1]
        try:
            with contextlib.redirect_stdout(StringIO()):
                result = target_fn(inp)
            if result == expected:
                passed += 1
            else:
                failures.append(f'  {inp!r} → got {result!r}, expected {expected!r}')
        except Exception as exc:
            failures.append(f'  {inp!r} → raised {type(exc).__name__}: {exc}')

    total      = len(test_cases)
    corr_score = round((passed / total) * 70) if total else 0
    score     += corr_score

    if passed == total:
        lines.append(f"✓ Correctness — passed all {total}/{total} test cases.")
    else:
        lines.append(f"✗ Correctness — passed {passed}/{total} test cases.")
        lines.extend(failures)

    qual_score, qual_notes = 0, []
    for node in ast.walk(tree):
        if isinstance(node, ast.FunctionDef):
            if (node.body and isinstance(node.body[0], ast.Expr)
                    and isinstance(getattr(node.body[0],"value",None), ast.Constant)
                    and isinstance(node.body[0].value.value, str)):
                qual_score += 5; qual_notes.append("docstring (+5)")
            if len(node.name) > 3 and node.name not in ("func","test"):
                qual_score += 5; qual_notes.append(f'name "{node.name}" (+5)')
            # Check type hints
            if node.returns or any(a.annotation for a in node.args.args):
                qual_score += 5; qual_notes.append("type hints (+5)")
            # Check for efficient implementation (collections usage)
            for child in ast.walk(node):
                if isinstance(child, ast.Import) or isinstance(child, ast.ImportFrom):
                    qual_score += 5; qual_notes.append("imports used (+5)"); break

    score = min(score + qual_score, 100)
    if qual_notes:
        lines.append("✓ Quality — " + ", ".join(dict.fromkeys(qual_notes)))
    else:
        lines.append("○ Quality — add docstring, type hints, and descriptive function name.")

    lines += [
        f"\n── Score ──────────────────────────",
        f"   Correctness : {corr_score:>3}/70",
        f"   Quality     : {min(qual_score,30):>3}/30",
        f"   Total       : {score:>3}/100",
    ]
    return {"task_score": score, "feedback": "\n".join(lines)}

def llm_mentor_feedback(code, task_desc, ast_fb, score, api_key) -> str:
    prompt = (
        f"Task: {task_desc}\nCode:\n```python\n{code[:1200]}\n```\n"
        f"Score: {score}/100\nNotes:\n{ast_fb}\n\n"
        "Write 3 sentences of mentor feedback: what they did well (be specific), "
        "one concrete improvement with example, overall impression. Under 90 words."
    )
    try:
        return _call_claude(api_key, prompt, max_tokens=250)
    except Exception as e:
        return f"*(Mentor feedback unavailable: {e})*"

# Final score calculation

def compute_final_score(match_score, task_score, github_signals, attempts, time_taken) -> dict:
    gh_bonus = 0
    if github_signals and not github_signals.get("error"):
        n_langs = len(github_signals.get("languages", {}))
        stars   = github_signals.get("total_stars", 0)
        pushes  = github_signals.get("push_events", 0)
        gh_bonus = min(10,
            (2 if n_langs >= 3 else 0) +
            (3 if stars >= 10 else 1 if stars >= 1 else 0) +
            (5 if pushes >= 10 else 2 if pushes >= 3 else 0))

    base  = int(0.55 * match_score + 0.35 * task_score + 0.10 * gh_bonus * 10)
    final = min(base, 100)

    if   attempts == 1 and time_taken < 90: behavior = "Fast Problem Solver ⚡"
    elif attempts > 3:                       behavior = "Persistent Learner 🔁"
    elif time_taken > 180:                   behavior = "Deep Thinker 🧠"
    else:                                    behavior = "Steady Performer 🎯"

    return {"final": final, "gh_bonus": gh_bonus, "behavior": behavior,
            "confidence": "HIGH" if final>75 else ("MEDIUM" if final>50 else "LOW")}

# Final analysis generator

def generate_final_analysis(
    candidate_name: str,
    job_title: str,
    match_score: int,
    evidence_chain: list,
    github_data: dict,
    task_eval: dict,
    mcq_results: list,
    short_answer_results: list,
    bias_passed: bool,
    final_score_data: dict,
    api_key: str,
) -> dict:
    """
    Synthesises ALL pipeline outputs into a structured final analysis report.
    Every sub-section is grounded in real signals — nothing invented.

    Returns:
    {
      "hire_recommendation": "STRONG YES | YES | MAYBE | NO",
      "executive_summary":   str,
      "signal_breakdown":    { skill_match, task, mcq, short_answer, github, bias },
      "top_strengths":       [ { skill, evidence, source } × 3 ],
      "critical_gaps":       [ { skill, severity, upskill_suggestion } × N ],
      "risk_flags":          [ str ],
      "next_steps":          [ str ],
      "overall_score":       int,
      "confidence":          str,
    }
    """
    verified  = [e for e in evidence_chain if e["status"] == "verified"]
    missing   = [e for e in evidence_chain if e["status"] == "missing"]
    top3_v    = [e["skill"] for e in verified[:3]]
    top3_m    = [e["skill"] for e in missing[:3]]

    mcq_score = 0
    if mcq_results:
        correct   = sum(1 for r in mcq_results if r.get("correct"))
        mcq_score = round(correct / len(mcq_results) * 100)

    sa_score = 0
    if short_answer_results:
        sa_score  = round(sum(r.get("score", 0) for r in short_answer_results)
                          / (len(short_answer_results) * 10) * 100)

    gh_langs    = list(github_data.get("languages", {}).keys())[:3]
    gh_stars    = github_data.get("total_stars", 0)
    gh_pushes   = github_data.get("push_events", 0)
    has_github  = bool(gh_langs)

    overall     = final_score_data.get("final", match_score)
    behavior    = final_score_data.get("behavior", "")
    confidence  = final_score_data.get("confidence", "MEDIUM")

    # Determine hire recommendation
    if overall >= 78:
        hire_rec = "STRONG YES"
    elif overall >= 62:
        hire_rec = "YES"
    elif overall >= 45:
        hire_rec = "MAYBE"
    else:
        hire_rec = "NO"

    # Get upskill suggestions from taxonomy for missing skills
    gap_items = []
    for e in missing[:3]:
        skill = e["skill"]
        # Find adjacent skills in taxonomy
        adjacents = []
        for cat, skills in taxonomy.get("categories", {}).items():
            if skill in skills:
                adjacents = [s for s in skills if s != skill][:2]
                break
        suggestion = f"Study {', '.join(adjacents)}" if adjacents else f"Take a course on {skill}"
        severity   = "HIGH" if skill in (taxonomy.get("job_role_skills", {}).get(job_title, []))[:3] else "MEDIUM"
        gap_items.append({"skill": skill, "severity": severity,
                          "upskill_suggestion": suggestion})

    # Risk flags
    risk_flags = []
    if not has_github:
        risk_flags.append("No GitHub activity — real-world coding artefacts unverified")
    if task_eval.get("task_score", 0) < 40:
        risk_flags.append("Low live coding score — practical implementation skills need verification")
    if len(missing) > len(verified):
        risk_flags.append(f"More skill gaps ({len(missing)}) than verified skills ({len(verified)})")
    if mcq_results and mcq_score < 50:
        risk_flags.append(f"MCQ score {mcq_score}% suggests conceptual knowledge gaps")

    # Next steps
    next_steps = []
    if hire_rec in ("STRONG YES", "YES"):
        next_steps.append("Proceed to technical panel interview with focus on system design")
        next_steps.append(f"Verify depth in top skills: {', '.join(top3_v[:2])}")
        if top3_m:
            next_steps.append(f"Discuss plan for bridging gaps in: {', '.join(top3_m[:2])}")
    elif hire_rec == "MAYBE":
        next_steps.append("Conduct additional technical assessment focusing on gap skills")
        next_steps.append(f"Request portfolio samples demonstrating: {', '.join(top3_m[:2])}")
        if has_github:
            next_steps.append("Schedule a code walkthrough of GitHub projects")
    else:
        next_steps.append("Candidate does not meet minimum requirements at this time")
        next_steps.append(f"Suggest upskilling path: {', '.join([g['upskill_suggestion'] for g in gap_items[:2]])}")

    result = {
        "hire_recommendation": hire_rec,
        "signal_breakdown": {
            "skill_match":    match_score,
            "coding_task":    task_eval.get("task_score", 0),
            "mcq":            mcq_score,
            "short_answer":   sa_score,
            "github_bonus":   final_score_data.get("gh_bonus", 0),
            "bias_passed":    bias_passed,
        },
        "top_strengths": [
            {"skill": e["skill"],
             "evidence": e["because"][0] if e["because"] else "Semantically matched",
             "source": e["sources"][0] if e["sources"] else "semantic"}
            for e in verified[:3]
        ],
        "critical_gaps":    gap_items,
        "risk_flags":       risk_flags,
        "next_steps":       next_steps,
        "overall_score":    overall,
        "confidence":       confidence,
        "behavior_profile": behavior,
    }

    # LLM executive summary (grounded in the structured data above)
    if api_key:
        prompt = (
            f"Write a 3-sentence executive hiring summary for "
            f"'{candidate_name or 'the candidate'}' for '{job_title or 'this role'}'.\n\n"
            f"Data:\n"
            f"- Overall score: {overall}/100 ({confidence} confidence)\n"
            f"- Hire recommendation: {hire_rec}\n"
            f"- Verified skills with evidence: {', '.join(top3_v)}\n"
            f"- Critical gaps: {', '.join(top3_m) or 'none'}\n"
            f"- Coding task: {task_eval.get('task_score',0)}/100\n"
            f"- MCQ: {mcq_score}%\n"
            f"- GitHub: {'active' if has_github else 'not provided'} "
            f"({', '.join(gh_langs)}, {gh_stars} stars)\n"
            f"- Behavioral profile: {behavior}\n\n"
            f"Be specific. Name skills. State recommendation clearly. Under 90 words."
        )
        try:
            result["executive_summary"] = _call_claude(api_key, prompt, max_tokens=200)
        except Exception:
            result["executive_summary"] = (
                f"{candidate_name or 'Candidate'} scored {overall}/100 with "
                f"{'strong' if hire_rec in ('STRONG YES','YES') else 'partial'} alignment to {job_title or 'the role'}. "
                f"Verified strengths in {', '.join(top3_v[:2]) or 'assessed skills'}. "
                f"{'Recommend proceeding to interview.' if hire_rec in ('STRONG YES','YES') else 'Additional assessment recommended.'}"
            )
    else:
        result["executive_summary"] = (
            f"{candidate_name or 'Candidate'} achieved {overall}/100 overall. "
            f"Key verified skills: {', '.join(top3_v[:3]) or 'none'}. "
            f"Gaps identified: {', '.join(top3_m[:3]) or 'none'}. "
            f"Recommendation: {hire_rec}."
        )

    return result

# Session state defaults

_defaults = {
    "start_time": None, "attempts": 0, "match_score": 0,
    "matched": [], "missing": [], "bg_skills": [], "c_skills": [],
    "current_task": None, "github_data": {}, "evidence_chain": [],
    "extraction_method": "dataset", "task_eval": {},
    "selected_job": None,
    # assessment state
    "resume_text": "",          # text extracted from uploaded resume file
    "resume_source": "",        # label of parser used
    "assessment": None,         # full multi-Q assessment dict
    "mcq_answers": {},          # {q_idx: selected_option}
    "mcq_submitted": False,
    "mcq_results": [],
    "short_answers": {},        # {q_idx: text}
    "sa_submitted": False,
    "sa_results": [],
    "final_analysis": None,     # complete final analysis dict
    "bias_passed": True,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.start_time is None:
    st.session_state.start_time = time.time()

# Sidebar

with st.sidebar:
    st.markdown("### 🔬 TalentLens")
    st.markdown("*Signal Intelligence · Impact Area 01*")
    st.markdown("---")

    api_key  = st.text_input("Anthropic API Key", type="password",
                              placeholder="sk-ant-…",
                              help="LLM skill NER, explainability report, dynamic tasks, mentor feedback")
    gh_token = st.text_input("GitHub Token (optional)", type="password",
                              placeholder="ghp_…",
                              help="Raises rate limit 60 → 5000 req/hr")

    if api_key:  st.success("✦ LLM features active")
    if gh_token: st.success("✦ GitHub auth active")

    st.markdown("---")

    # Dataset status panel
    st.markdown("**📦 Dataset Status**")
    taxonomy_ok  = bool(taxonomy.get("categories"))
    resume_ok    = (DATA_DIR / "Resume.csv").exists()
    postings_ok  = not jobs_df.empty

    st.markdown(f"{'✅' if taxonomy_ok  else '⚠️'} Lightcast Skills Taxonomy `({sum(len(v) for v in taxonomy.get('categories',{}).values())} skills)`")
    st.markdown(f"{'✅' if resume_ok    else '📥'} Kaggle Resume Dataset `({'loaded' if resume_ok else 'place Resume.csv in data/'})`")
    st.markdown(f"{'✅' if postings_ok  else '📥'} LinkedIn Job Postings `({len(jobs_df)} postings loaded)`")

    if not resume_ok or not postings_ok:
        with st.expander("How to load full datasets"):
            st.markdown("""
1. **Resume Dataset** → [Kaggle](https://www.kaggle.com/datasets/snehaanbhawal/resume-dataset) → download `Resume.csv` → place in `data/`
2. **Job Postings** → [Kaggle](https://www.kaggle.com/datasets/arshkon/linkedin-job-postings) → download `postings.csv` → place in `data/`
            """)

    st.markdown("---")
    st.markdown(f"**Embedding:** `all-MiniLM-L6-v2`")
    st.markdown(f"**LLM:** `{CLAUDE_MODEL}`")
    st.markdown(f"**Skill vocab:** `{len(skill_vocab):,}` terms from datasets")

# Hero header

vocab_count = sum(len(v) for v in taxonomy.get("categories", {}).values())
st.markdown(f"""
<div class="hero">
  <div class="hero-eyebrow">🔬 Techkriti '26 × Eightfold AI · Impact Area 01 · Signal Extraction & Verification</div>
  <h1 class="hero-title">TalentLens: <span>Signal Intelligence</span></h1>
  <p class="hero-sub">Defeating AI-spam applications by verifying genuine capability from GitHub artefacts, real-world skill signals, and live coding — grounded in {vocab_count}+ skills from Kaggle Resume Dataset & Lightcast Open Skills taxonomy.</p>
  <div class="hero-chips">
    <span class="hero-chip">🐙 GitHub Parsing</span>
    <span class="hero-chip">📦 Kaggle Resume Dataset</span>
    <span class="hero-chip">💼 LinkedIn Job Postings</span>
    <span class="hero-chip">🧠 Lightcast Taxonomy</span>
    <span class="hero-chip">🔗 Semantic Matching</span>
    <span class="hero-chip">📋 Evidence Chain</span>
    <span class="hero-chip">🛡️ Bias Audit</span>
  </div>
</div>
""", unsafe_allow_html=True)

# Section A: job requirements

st.markdown('<p class="sec-label">Section A · Role Definition</p>', unsafe_allow_html=True)
st.markdown(
    '<p class="sec-title">Job Requirements '
    '<span class="dataset-badge">📦 LinkedIn Job Postings</span></p>',
    unsafe_allow_html=True
)
st.markdown('<p class="sec-desc">Search the LinkedIn Job Postings dataset to auto-populate requirements, or enter manually. The judge can change this live during the demo.</p>', unsafe_allow_html=True)

a1, a2 = st.columns([1, 2], gap="large")
with a1:
    job_title = st.text_input("Role Title", placeholder="e.g. Machine Learning Engineer")
with a2:
    # Live job search from dataset
    search_query = st.text_input("🔍 Search Job Postings Dataset", placeholder="e.g. NLP, backend, data science…")

if search_query:
    job_results = search_job_postings(search_query, max_results=6)
    if job_results:
        st.markdown("**Select a job to auto-populate requirements:**")
        cols = st.columns(min(len(job_results), 3), gap="medium")
        for i, job in enumerate(job_results[:6]):
            with cols[i % 3]:
                if st.button(f"📋 {job['title']}\n{job['company']}", key=f"job_{i}"):
                    st.session_state.selected_job = job
                    st.rerun()
    else:
        st.info("No matching postings found. Try a different query.")

if st.session_state.selected_job:
    sel = st.session_state.selected_job
    st.markdown(f"""
    <div class="job-card" style="border-color:#4f46e5;cursor:default">
      <div class="job-card-title">✅ Selected: {sel['title']} @ {sel['company']}</div>
      <div class="job-card-meta">📍 {sel.get('location','') or 'Remote'} · From LinkedIn Job Postings Dataset</div>
      <div style="margin-top:.5rem;font-size:.82rem;color:#374151">{sel.get('desc_preview','')}</div>
    </div>""", unsafe_allow_html=True)
    prefilled_skills = ", ".join(sel["skills"])
    if st.button("✕ Clear selection"):
        st.session_state.selected_job = None
        st.rerun()
else:
    prefilled_skills = ""

job_desc = st.text_area(
    "Required Skills (comma-separated)",
    value=prefilled_skills,
    height=80,
    placeholder="Auto-populated from dataset, or type manually: Python, PyTorch, Docker…",
)

st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

# Section B: candidate signals

st.markdown('<p class="sec-label">Section B · Candidate Signal Sources</p>', unsafe_allow_html=True)
st.markdown(
    '<p class="sec-title">Multi-Source Signal Collection '
    '<span class="dataset-badge">📦 Resume Dataset NER</span>'
    '<span class="llm-badge">✦ LLM</span></p>',
    unsafe_allow_html=True
)
st.markdown('<p class="sec-desc">Skills extracted from resume text using vocabulary derived from 2,400+ real resumes in the Kaggle Resume Dataset — no hardcoded keyword lists.</p>', unsafe_allow_html=True)

b1, b2 = st.columns(2, gap="large")
with b1:
    candidate_name   = st.text_input("Candidate Name", placeholder="Alex Kumar (for report only)")
    github_username  = st.text_input("🐙 GitHub Username", placeholder="e.g. torvalds")
    candidate_skills = st.text_area("Declared Skills (comma-separated)", height=80,
                                    placeholder="Python, PyTorch, Docker, PostgreSQL…")
with b2:
    # Resume file upload (PDF / DOCX / JPG / PNG)
    st.markdown(
        "**📄 Upload Resume** "
        "<small style='color:#6b7280'>(PDF, DOCX, JPG, PNG)</small>",
        unsafe_allow_html=True
    )
    uploaded_resume = st.file_uploader(
        "resume_uploader", type=["pdf","docx","jpg","jpeg","png"],
        label_visibility="collapsed",
        help="Upload a resume. Text is extracted automatically and fed into the NER pipeline."
    )

    # Parse on upload
    if uploaded_resume is not None:
        if (st.session_state.resume_source == "" or
                uploaded_resume.name not in st.session_state.resume_source):
            with st.spinner("Parsing resume…"):
                rtext, rlabel = parse_resume_file(uploaded_resume, api_key)
            st.session_state.resume_text   = rtext
            st.session_state.resume_source = f"{uploaded_resume.name} ({rlabel})"

    # Show parsed text preview + auto-fill
    if st.session_state.resume_text:
        preview = st.session_state.resume_text[:300].replace("\n"," ")
        st.markdown(f"""
        <div style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:10px;
                    padding:.75rem 1rem;margin-top:.4rem">
          <div style="display:flex;align-items:center;gap:.5rem;margin-bottom:.4rem">
            <span style="color:#16a34a;font-size:.8rem;font-weight:700">✅ Parsed</span>
            <span style="color:#6b7280;font-size:.75rem">{st.session_state.resume_source}</span>
          </div>
          <p style="font-size:.78rem;color:#374151;margin:0;line-height:1.5">
            {preview}{"…" if len(st.session_state.resume_text) > 300 else ""}
          </p>
        </div>""", unsafe_allow_html=True)

        # Auto-fill declared skills if field is empty
        if not candidate_skills.strip() and api_key:
            if st.button("✨ Auto-fill skills from resume"):
                with st.spinner("Extracting skills…"):
                    auto_skills = llm_extract_skills(st.session_state.resume_text, api_key)
                if auto_skills:
                    st.session_state["_autofill_skills"] = ", ".join(auto_skills[:15])
                    st.rerun()
    else:
        st.markdown(
            '<p style="font-size:.78rem;color:#9ca3af;margin-top:.4rem">'
            'Upload a resume to auto-extract skills into the pipeline.</p>',
            unsafe_allow_html=True
        )

    # Apply autofill if triggered
    if "_autofill_skills" in st.session_state:
        candidate_skills = st.session_state.pop("_autofill_skills")

    linkedin_input = st.text_area("🔗 LinkedIn Bio/Experience", height=68,
                                   placeholder="Paste LinkedIn summary…")

st.markdown("**Demographic fields for bias audit** *(should never affect score)*")
dc1, dc2, dc3 = st.columns(3, gap="medium")
with dc1: dem_name = st.text_input("Name / Gender signal", placeholder="e.g. Priya Sharma")
with dc2: dem_univ = st.text_input("University",           placeholder="e.g. IIT Kanpur")
with dc3: dem_year = st.text_input("Graduation Year",      placeholder="e.g. 2023")

run_btn = st.button("🔬 Run Full Signal Analysis →")

# Run analysis pipeline

if run_btn:
    if not job_desc:
        st.warning("Please enter job requirements or select a job from the dataset.")
        st.stop()

    j_skills = [s.strip() for s in job_desc.split(",") if s.strip()]
    c_skills = [s.strip() for s in candidate_skills.split(",") if s.strip()]
    st.session_state.c_skills = c_skills

    progress = st.progress(0, text="Starting pipeline…")

    # GitHub profile
    gh_data = {}
    if github_username.strip():
        progress.progress(10, text=f"🐙 Fetching GitHub @{github_username}…")
        gh_data = parse_github_profile(github_username.strip(), gh_token.strip() if gh_token else "")
        if gh_data.get("error"):
            st.warning(f"GitHub: {gh_data['error']}")
    st.session_state.github_data = gh_data

    # Skill extraction from resume + LinkedIn + GitHub bio
    resume_text_src = st.session_state.resume_text  # from file upload
    bg_text = " ".join(filter(None, [resume_text_src, linkedin_input, gh_data.get("bio","")]))
    bg_skills = []
    if bg_text.strip():
        if api_key:
            progress.progress(30, text="🧠 Extracting skills via Claude NER (grounded in Lightcast taxonomy)…")
            bg_skills = llm_extract_skills(bg_text, api_key)
            st.session_state.extraction_method = "llm+dataset"
        else:
            progress.progress(30, text="📦 Extracting skills using Resume Dataset vocabulary…")
            bg_skills = dataset_extract_skills(bg_text)
            st.session_state.extraction_method = "dataset"

    gh_skill_names  = [s["skill"] for s in gh_data.get("inferred_skills", [])]
    all_c_skills    = list(dict.fromkeys(c_skills + bg_skills + gh_skill_names))
    st.session_state.bg_skills = bg_skills

    # Semantic skill matching
    progress.progress(50, text="🔗 Semantic matching with all-MiniLM-L6-v2…")
    if j_skills and all_c_skills:
        result = match_skills(j_skills, all_c_skills)
        st.session_state.match_score = result["score"]
        st.session_state.matched     = result["matched_skills"]
        st.session_state.missing     = result["missing_skills"]
    else:
        st.session_state.match_score = 0
        st.session_state.matched     = []
        st.session_state.missing     = j_skills

    # Generate full assessment
    progress.progress(68, text="💻 Generating multi-format assessment grounded in job dataset…")
    if api_key and st.session_state.matched:
        st.session_state.assessment = llm_generate_assessment(
            st.session_state.matched,
            st.session_state.missing,
            job_title or "",
            api_key,
        )
    else:
        st.session_state.assessment = {
            "coding_task":  _default_task(),
            "mcq":          [],
            "short_answer": [],
            "total_marks":  100,
        }
    st.session_state.current_task    = st.session_state.assessment["coding_task"]
    st.session_state.mcq_answers     = {}
    st.session_state.mcq_submitted   = False
    st.session_state.mcq_results     = []
    st.session_state.short_answers   = {}
    st.session_state.sa_submitted    = False
    st.session_state.sa_results      = []
    st.session_state.task_eval       = {}
    st.session_state.final_analysis  = None

    # Build evidence chain
    progress.progress(85, text="📋 Building evidence chain…")
    st.session_state.evidence_chain = build_evidence_chain(
        st.session_state.matched, st.session_state.missing,
        gh_data, bg_skills, {},
    )
    st.session_state.attempts    = 0
    st.session_state.start_time  = time.time()

    progress.progress(100, text="✅ Analysis complete.")
    time.sleep(0.3)
    progress.empty()

has_results = bool(st.session_state.matched or st.session_state.missing)

if has_results:

    # GITHUB PANEL
    gh = st.session_state.github_data
    if gh and not gh.get("error") and gh.get("languages"):
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<p class="sec-label">GitHub Signal Report</p>', unsafe_allow_html=True)
        st.markdown('<p class="sec-title">🐙 Repository & Language Analysis</p>', unsafe_allow_html=True)

        c1,c2,c3,c4 = st.columns(4, gap="medium")
        for col, val, lab in [
            (c1, gh.get("total_repos",0), "Public Repos"),
            (c2, gh.get("total_stars",0), "Total Stars ⭐"),
            (c3, gh.get("push_events",0), "Recent Pushes"),
            (c4, len(gh.get("languages",{})), "Languages"),
        ]:
            with col:
                st.markdown(f'<div class="stat-card"><div class="stat-val">{val}</div><div class="stat-lab">{lab}</div></div>', unsafe_allow_html=True)

        langs = gh.get("languages", {})
        if langs:
            total_pct = sum(langs.values()) or 1
            segs = "".join(
                f'<div class="lang-seg" style="width:{p/total_pct*100:.1f}%;background:{LANG_COLORS.get(l,LANG_COLORS["default"])}"></div>'
                for l,p in list(langs.items())[:10]
            )
            dots = "".join(
                f'<span class="lang-dot"><span style="width:10px;height:10px;border-radius:50%;background:{LANG_COLORS.get(l,LANG_COLORS["default"])};display:inline-block"></span>{l} {p:.0f}%</span>'
                for l,p in list(langs.items())[:8]
            )
            st.markdown(f"""
            <div style="margin-top:1rem">
              <p style="font-size:.82rem;font-weight:600;color:#374151;margin-bottom:.4rem">
                Language Distribution
                <span class="dataset-badge">📦 Lightcast taxonomy mapping</span>
              </p>
              <div class="lang-bar-bg">{segs}</div>
              <div class="lang-legend">{dots}</div>
            </div>""", unsafe_allow_html=True)

        if gh.get("top_repos"):
            st.markdown('<p style="font-size:.82rem;font-weight:600;color:#374151;margin:1.2rem 0 .5rem 0">Top Repositories</p>', unsafe_allow_html=True)
            repo_cols = st.columns(min(len(gh["top_repos"]),3), gap="medium")
            for i, repo in enumerate(gh["top_repos"][:3]):
                with repo_cols[i]:
                    topics_html = "".join(f'<span style="background:#388bfd26;color:#58a6ff;border:1px solid #1f6feb;font-size:.7rem;padding:.1rem .45rem;border-radius:12px;margin:.1rem">{t}</span>' for t in repo.get("topics",[])[:3])
                    st.markdown(f"""
                    <div style="background:#0d1117;border:1px solid #30363d;border-radius:10px;padding:1rem">
                      <div style="color:#58a6ff;font-size:.85rem;font-weight:600;margin-bottom:.3rem">{repo['name']} <span style="color:#f0e68c">⭐ {repo['stars']}</span></div>
                      <p style="color:#8b949e;font-size:.78rem;margin:0 0 .5rem 0">{repo.get('description','') or 'No description'}</p>
                      <span style="color:#f0883e;font-size:.75rem">● {repo['language']}</span>
                      <div style="margin-top:.4rem;display:flex;flex-wrap:wrap;gap:.25rem">{topics_html}</div>
                    </div>""", unsafe_allow_html=True)

        if gh.get("inferred_skills"):
            st.markdown(f'<p style="font-size:.82rem;font-weight:600;color:#374151;margin:1.2rem 0 .4rem 0">Inferred skills <span class="dataset-badge">📦 Lightcast language→skill map</span></p>', unsafe_allow_html=True)
            pills = '<div class="pill-row">' + "".join(
                f'<span class="pill github-p">🐙 {s["skill"]} <span style="opacity:.5;font-size:.72rem">({s["strength"]})</span></span>'
                for s in gh["inferred_skills"][:12]
            ) + '</div>'
            st.markdown(pills, unsafe_allow_html=True)

    # SKILL MATCH
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    em = st.session_state.extraction_method
    em_label = ("LLM + Dataset" if em == "llm+dataset" else "Kaggle Resume Dataset vocabulary")
    st.markdown(f'<p class="sec-label">Semantic Skill Match · extraction via {em_label}</p>', unsafe_allow_html=True)

    score   = st.session_state.match_score
    matched = st.session_state.matched
    missing = st.session_state.missing

    r = 40; c = round(2*3.14159*r,2)
    dash = round(c*score/100,2); gap = round(c-dash,2)
    rc = "#16a34a" if score>=75 else ("#ca8a04" if score>=50 else "#dc2626")
    cc = "high" if score>=75 else ("medium" if score>=50 else "low")
    cl = "Strong" if score>=75 else ("Partial" if score>=50 else "Weak")

    sm1,sm2,sm3 = st.columns([1,2,2], gap="large")
    with sm1:
        st.markdown(f"""
        <div style="display:flex;flex-direction:column;align-items:center;gap:.6rem;padding-top:.3rem">
          <div class="ring-wrap">
            <svg width="100" height="100" viewBox="0 0 100 100">
              <circle cx="50" cy="50" r="{r}" fill="none" stroke="#e5e7eb" stroke-width="9"/>
              <circle cx="50" cy="50" r="{r}" fill="none" stroke="{rc}" stroke-width="9" stroke-dasharray="{dash} {gap}" stroke-linecap="round"/>
            </svg>
            <div class="ring-inner"><span class="ring-num">{score}</span><span class="ring-denom">/ 100</span></div>
          </div>
          <span class="badge {cc}">{cl} Match</span>
        </div>""", unsafe_allow_html=True)
    with sm2:
        st.markdown("**✅ Matched Skills**")
        st.markdown('<div class="pill-row">' + ("".join(
            f'<span class="pill matched">✓ {m["job_skill"]} <span style="opacity:.55;font-size:.72rem">{int(m["similarity"]*100)}%</span></span>'
            for m in matched
        ) or '<span style="color:#9ca3af;font-size:.86rem">No matches found.</span>') + '</div>', unsafe_allow_html=True)
    with sm3:
        st.markdown("**⚠️ Missing Skills**")
        st.markdown('<div class="pill-row">' + ("".join(
            f'<span class="pill missing">✗ {s}</span>' for s in missing
        ) or '<span style="color:#16a34a;font-size:.86rem;font-weight:600">All skills covered! 🎉</span>') + '</div>', unsafe_allow_html=True)

    if st.session_state.bg_skills:
        st.markdown(f'<p style="font-size:.82rem;font-weight:600;color:#374151;margin:1.2rem 0 .4rem 0">Extracted background skills <span class="dataset-badge">📦 Resume Dataset vocab</span></p>', unsafe_allow_html=True)
        verified = {m["candidate_skill"] for m in matched}
        st.markdown('<div class="pill-row">' + "".join(
            f'<span class="pill verified">● {s} verified</span>' if s in verified
            else f'<span class="pill dataset-p">○ {s}</span>'
            for s in st.session_state.bg_skills
        ) + '</div>', unsafe_allow_html=True)

    # Evidence chain builder
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="sec-label">Explainability · Core Requirement of Impact Area 01</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">📋 Evidence Chain — "We believe this candidate has skill X because…"</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-desc">Every claim is traceable to a real-world artefact. No black-box percentages.</p>', unsafe_allow_html=True)

    src_icons = {"github":"🐙","resume":"📄","semantic":"🔗","task":"💻","dataset":"📦"}
    for ev in st.session_state.evidence_chain:
        is_v      = ev["status"] == "verified"
        conf      = ev.get("confidence", "MEDIUM")
        cc2       = {"HIGH":"#16a34a","MEDIUM":"#ca8a04","LOW":"#dc2626"}.get(conf,"#ca8a04")
        cc2bg     = {"HIGH":"#dcfce7","MEDIUM":"#fef9c3","LOW":"#fee2e2"}.get(conf,"#fef9c3")
        border_c  = "#16a34a" if is_v else "#dc2626"
        status_ic = "✅" if is_v else "❌"

        src_tags = "".join(
            f'<span class="ev-src {s}">{src_icons.get(s,"📎")} {s}</span>'
            for s in ev.get("sources", [])
        )
        # Build bullet list with proper HTML escaping so raw tags never leak
        bullets = "".join(
            f'<div style="display:flex;gap:.4rem;margin-top:.25rem">' +
            f'<span style="color:{cc2};font-weight:700;flex-shrink:0;margin-top:.05rem">•</span>' +
            f'<span style="font-size:.82rem;color:#555">{html.escape(str(b))}</span></div>'
            for b in ev.get("because", [])
        )
        skill_esc = html.escape(ev["skill"])

        st.markdown(
            f'<div class="evidence-item" style="border-left:3px solid {border_c}">' +
            f'  <div style="font-size:1rem;flex-shrink:0;margin-top:.1rem">{status_ic}</div>' +
            f'  <div style="flex:1;min-width:0">' +
            f'    <div style="display:flex;align-items:center;gap:.5rem;flex-wrap:wrap;margin-bottom:.25rem">' +
            f'      <span class="evidence-claim">{skill_esc}</span>' +
            f'      <span style="font-size:.72rem;font-weight:700;color:{cc2};background:{cc2bg};padding:.1rem .45rem;border-radius:20px">{conf}</span>' +
            f'      {src_tags}' +
            f'    </div>' +
            f'    <div>{bullets}</div>' +
            f'  </div>' +
            f'</div>',
            unsafe_allow_html=True
        )

    if api_key and st.session_state.evidence_chain:
        with st.spinner("🧠 Synthesising report…"):
            narrative = llm_generate_explanation(
                st.session_state.evidence_chain, candidate_name,
                job_title or (st.session_state.selected_job or {}).get("title",""), api_key
            )
        st.markdown(f"""
        <div class="explain-card" style="margin-top:1rem">
          <div style="display:flex;align-items:center;gap:.6rem;margin-bottom:.75rem">
            <span style="font-size:1.1rem">🧠</span>
            <span style="font-weight:700;font-size:.9rem;color:#92400e">AI Intelligence Report</span>
            <span class="llm-badge">✦ LLM</span>
          </div>
          <p style="font-size:.92rem;color:#374151;line-height:1.8;margin:0">{narrative}</p>
        </div>""", unsafe_allow_html=True)

    # Bias audit
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="sec-label">Fairness · Bias Audit</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">🛡️ Score Invariance Proof</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-desc">Demonstrating the score is mathematically invariant to demographic signals. Required for any AI used in hiring decisions.</p>', unsafe_allow_html=True)

    dem_fields = {k:v for k,v in {"Name":dem_name,"University":dem_univ,"Grad Year":dem_year}.items() if v.strip()}
    j_s = [s.strip() for s in job_desc.split(",") if s.strip()]

    if dem_fields and (st.session_state.c_skills + st.session_state.bg_skills):
        audit = run_bias_audit(j_s, st.session_state.c_skills + st.session_state.bg_skills, dem_fields)
        ap = audit["all_pass"]
        st.markdown(f"""
        <div class="bias-card">
          <div style="display:flex;align-items:center;gap:.75rem;margin-bottom:1rem">
            <span style="font-size:1.3rem">{"✅" if ap else "⚠️"}</span>
            <span style="font-weight:700;color:{"#15803d" if ap else "#dc2626"}">
              {"BIAS-FREE — score invariant to all demographic fields" if ap else "WARNING — score variance detected"}
            </span>
          </div>
          <div style="background:white;border:1px solid #e5e7eb;border-radius:10px;padding:1rem">""",
        unsafe_allow_html=True)
        for ck in audit["checks"]:
            d = ck["delta"]; ps = "bias-pass" if ck["pass"] else "bias-fail"
            pass_color = "#16a34a" if ck["pass"] else "#dc2626"
            st.markdown(f"""
            <div style="display:flex;align-items:center;gap:1rem;padding:.65rem 0;border-bottom:1px solid #f3f4f6;font-size:.86rem">
              <span style="font-weight:600;color:#374151;min-width:130px">{ck['field']}</span>
              <span style="color:#6b7280;font-family:'DM Mono',monospace;font-size:.8rem;flex:1">"{ck['value']}"</span>
              <span style="color:#6b7280;font-size:.82rem">with: <strong>{ck['score_with']}</strong> → without: <strong>{ck['score_without']}</strong></span>
              <span style="font-weight:700;color:{pass_color}">Δ{d} {"✓" if ck["pass"] else "✗"}</span>
            </div>""", unsafe_allow_html=True)
        st.markdown('</div></div>', unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="bias-card">
          <p style="color:#374151;font-size:.88rem;margin:0">
            ✅ <strong>By design:</strong> scoring uses zero demographic features — only skill artefacts.
            Fill in the demographic fields above to run a live invariance proof during the demo.
          </p>
        </div>""", unsafe_allow_html=True)

    # LIVE MULTI-FORMAT ASSESSMENT
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="sec-label">Section C · Live Capability Verification</p>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sec-title">🧪 Multi-Format Assessment '
        f'{"<span class=\\'llm-badge\\'>✦ LLM Generated</span>" if api_key and st.session_state.matched else ""}'
        '<span class="dataset-badge">📦 Role-grounded via Job Dataset</span></p>',
        unsafe_allow_html=True
    )
    st.markdown('<p class="sec-desc">Three assessment formats — coding task, conceptual MCQ, and short-answer — covering both verified strengths and identified skill gaps.</p>', unsafe_allow_html=True)

    assessment = st.session_state.assessment or {"coding_task": _default_task(), "mcq": [], "short_answer": [], "total_marks": 100}
    task       = assessment.get("coding_task") or _default_task()
    mcqs       = assessment.get("mcq", [])
    sas        = assessment.get("short_answer", [])
    test_cases = task.get("test_cases", [])

    # PART 1: CODING TASK
    st.markdown("""
    <div style="display:flex;align-items:center;gap:.75rem;margin-bottom:.75rem">
      <div style="background:#4f46e5;color:white;border-radius:8px;padding:.3rem .65rem;font-weight:700;font-size:.8rem">Part 1</div>
      <span style="font-weight:700;color:#111827">Coding Task</span>
      <span style="font-size:.78rem;color:#6b7280">· 50 marks</span>
    </div>""", unsafe_allow_html=True)

    diff_str = task.get("difficulty","Easy")
    diff_cls  = diff_str.lower()
    diff_col  = {"easy":"#16a34a","medium":"#ca8a04","hard":"#dc2626"}.get(diff_cls,"#16a34a")
    diff_bg   = {"easy":"#dcfce7","medium":"#fef9c3","hard":"#fee2e2"}.get(diff_cls,"#dcfce7")
    tag_html  = "".join(f'<span style="background:#eff6ff;color:#1d4ed8;border:1px solid #bfdbfe;font-size:.76rem;padding:.15rem .5rem;border-radius:12px">{t}</span>' for t in task.get("tags",[]))
    hint_html = (f'<p style="font-size:.78rem;color:#6b7280;margin:.5rem 0 0 0">💡 <code style="background:#e0e7ff;color:#3730a3;padding:.1rem .4rem;border-radius:4px">{task.get("evaluation_hint","")}</code></p>' if task.get("evaluation_hint") else "")

    if api_key and st.session_state.matched:
        rg1, _ = st.columns([1,5])
        with rg1:
            if st.button("↻ Regenerate All"):
                with st.spinner("Regenerating assessment…"):
                    st.session_state.assessment = llm_generate_assessment(
                        st.session_state.matched, st.session_state.missing,
                        job_title or "", api_key
                    )
                    st.session_state.current_task  = st.session_state.assessment["coding_task"]
                    st.session_state.mcq_answers   = {}
                    st.session_state.mcq_submitted = False
                    st.session_state.mcq_results   = []
                    st.session_state.short_answers = {}
                    st.session_state.sa_submitted  = False
                    st.session_state.sa_results    = []
                    st.session_state.task_eval     = {}
                    st.session_state.final_analysis= None
                    st.session_state.attempts      = 0
                    st.session_state.start_time    = time.time()
                st.rerun()

    st.markdown(f"""
    <div class="task-card">
      <div style="display:flex;align-items:flex-start;gap:1rem">
        <div style="background:#4f46e5;color:white;border-radius:8px;width:36px;height:36px;display:flex;align-items:center;justify-content:center;font-size:1rem;flex-shrink:0">λ</div>
        <div style="flex:1">
          <div style="display:flex;align-items:center;gap:.6rem;margin-bottom:.3rem;flex-wrap:wrap">
            <span style="font-weight:700;font-size:.95rem;color:#111827">{task["title"]}</span>
            <span style="background:{diff_bg};color:{diff_col};font-size:.74rem;font-weight:700;padding:.15rem .55rem;border-radius:20px">{diff_str}</span>
          </div>
          <p style="font-size:.875rem;color:#4b5563;margin:0 0 .6rem 0">{task["description"]}</p>
          <div style="display:flex;gap:.4rem;flex-wrap:wrap">{tag_html}</div>
          {hint_html}
        </div>
      </div>
    </div>""", unsafe_allow_html=True)

    user_code = st.text_area("Your Code", height=160,
                              placeholder=task.get("placeholder","def solution(...):\n    pass"),
                              key="code_input")

    if st.button("▶ Submit Code"):
        st.session_state.attempts += 1
        if user_code.strip():
            with st.spinner("Evaluating code…"):
                tc_to_use  = test_cases if test_cases else _default_task()["test_cases"]
                eval_r     = evaluate_code(user_code, tc_to_use)
                eval_r["tags"] = task.get("tags", [])
                st.session_state.task_eval  = eval_r
                time_taken = round(time.time() - st.session_state.start_time, 2)
                eval_r["time_taken"] = time_taken
                # Rebuild evidence chain with task result
                st.session_state.evidence_chain = build_evidence_chain(
                    st.session_state.matched, st.session_state.missing,
                    st.session_state.github_data, st.session_state.bg_skills, eval_r,
                )
            task_score = eval_r["task_score"]
            tc_col = "#16a34a" if task_score>=70 else ("#ca8a04" if task_score>=40 else "#dc2626")
            tc_lab = "Excellent" if task_score>=70 else ("Acceptable" if task_score>=40 else "Needs Work")

            r1,r2,r3 = st.columns(3, gap="medium")
            with r1: st.markdown(f'<div class="stat-card"><div class="stat-val" style="color:{tc_col}">{task_score}<span style="font-size:.95rem;color:#9ca3af">/100</span></div><div class="stat-lab">Code Score · {tc_lab}</div></div>', unsafe_allow_html=True)
            with r2: st.markdown(f'<div class="stat-card"><div class="stat-val">{time_taken}s</div><div class="stat-lab">Time Taken</div></div>', unsafe_allow_html=True)
            with r3: st.markdown(f'<div class="stat-card"><div class="stat-val">{st.session_state.attempts}</div><div class="stat-lab">Attempt{"s" if st.session_state.attempts!=1 else ""}</div></div>', unsafe_allow_html=True)

            st.markdown(f'<div class="feedback-mono">{eval_r["feedback"]}</div>', unsafe_allow_html=True)

            if api_key:
                with st.spinner("🧠 Generating mentor feedback…"):
                    mfb = llm_mentor_feedback(user_code, task["description"], eval_r["feedback"], task_score, api_key)
                st.markdown(f'<div class="mentor-block" style="margin-top:.75rem">🎓 {mfb}</div>', unsafe_allow_html=True)
        else:
            st.warning("Please write your code solution.")

    # PART 2: MCQ
    if mcqs:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown("""
        <div style="display:flex;align-items:center;gap:.75rem;margin-bottom:.75rem">
          <div style="background:#7c3aed;color:white;border-radius:8px;padding:.3rem .65rem;font-weight:700;font-size:.8rem">Part 2</div>
          <span style="font-weight:700;color:#111827">Conceptual MCQ</span>
          <span style="font-size:.78rem;color:#6b7280">· 30 marks · 4 questions · one correct answer each</span>
        </div>""", unsafe_allow_html=True)

        for qi, q in enumerate(mcqs):
            skill_tag = f'<span class="dataset-badge">📦 {q.get("skill","")}</span>' if q.get("skill") else ""
            st.markdown(f"""
            <div style="background:#faf5ff;border:1px solid #ddd6fe;border-radius:12px;padding:1rem 1.25rem;margin-bottom:.75rem">
              <p style="font-weight:600;font-size:.9rem;color:#111827;margin:0 0 .75rem 0">
                Q{qi+1}. {q['question']} {skill_tag}
              </p>""", unsafe_allow_html=True)

            if not st.session_state.mcq_submitted:
                options = q.get("options", [])
                chosen  = st.radio(
                    f"q{qi}", options, index=None,
                    key=f"mcq_{qi}", label_visibility="collapsed"
                )
                if chosen:
                    letter = chosen[0]  # "A", "B", "C", "D"
                    st.session_state.mcq_answers[qi] = letter
            else:
                # Show result after submission
                res = st.session_state.mcq_results[qi] if qi < len(st.session_state.mcq_results) else {}
                for opt in q.get("options", []):
                    letter = opt[0]
                    correct_letter = q.get("answer","A")
                    chosen_letter  = st.session_state.mcq_answers.get(qi,"")
                    if letter == correct_letter:
                        st.markdown(f'<p style="background:#dcfce7;border-radius:6px;padding:.3rem .75rem;font-size:.86rem;color:#15803d">✅ {opt}</p>', unsafe_allow_html=True)
                    elif letter == chosen_letter and letter != correct_letter:
                        st.markdown(f'<p style="background:#fee2e2;border-radius:6px;padding:.3rem .75rem;font-size:.86rem;color:#dc2626">✗ {opt}</p>', unsafe_allow_html=True)
                    else:
                        st.markdown(f'<p style="font-size:.86rem;color:#6b7280;padding:.3rem .75rem;margin:0">{opt}</p>', unsafe_allow_html=True)
                if q.get("explanation"):
                    st.markdown(f'<p style="font-size:.8rem;color:#7c3aed;margin:.25rem 0 0 0">💡 {q["explanation"]}</p>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        if not st.session_state.mcq_submitted:
            if st.button("Submit MCQ Answers →"):
                results = []
                for qi, q in enumerate(mcqs):
                    chosen  = st.session_state.mcq_answers.get(qi, "")
                    correct = q.get("answer","A")
                    results.append({"correct": chosen == correct,
                                    "chosen": chosen, "correct_answer": correct})
                st.session_state.mcq_results   = results
                st.session_state.mcq_submitted = True
                score_pct = round(sum(1 for r in results if r["correct"]) / len(results) * 100)
                st.success(f"MCQ submitted — {score_pct}% correct ({sum(1 for r in results if r['correct'])}/{len(results)})")
                st.rerun()
        else:
            correct_count = sum(1 for r in st.session_state.mcq_results if r.get("correct"))
            score_pct     = round(correct_count / len(mcqs) * 100)
            col_c = "#16a34a" if score_pct>=70 else ("#ca8a04" if score_pct>=50 else "#dc2626")
            st.markdown(f'<div class="stat-card" style="display:inline-block;margin-top:.5rem"><div class="stat-val" style="color:{col_c}">{score_pct}%</div><div class="stat-lab">MCQ Score · {correct_count}/{len(mcqs)} correct</div></div>', unsafe_allow_html=True)

    # PART 3: SHORT ANSWER
    if sas:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown("""
        <div style="display:flex;align-items:center;gap:.75rem;margin-bottom:.75rem">
          <div style="background:#0f766e;color:white;border-radius:8px;padding:.3rem .65rem;font-weight:700;font-size:.8rem">Part 3</div>
          <span style="font-weight:700;color:#111827">Short Answer</span>
          <span style="font-size:.78rem;color:#6b7280">· 20 marks · write 2-4 sentences per answer</span>
        </div>""", unsafe_allow_html=True)

        for si, q in enumerate(sas):
            skill_tag = f'<span class="dataset-badge">📦 {q.get("skill","")}</span>' if q.get("skill") else ""
            st.markdown(f"""
            <div style="background:#f0fdfa;border:1px solid #99f6e4;border-radius:12px;padding:1rem 1.25rem;margin-bottom:.75rem">
              <p style="font-weight:600;font-size:.9rem;color:#111827;margin:0 0 .6rem 0">
                SA{si+1}. {q['question']} {skill_tag}
              </p>""", unsafe_allow_html=True)

            if not st.session_state.sa_submitted:
                ans = st.text_area(
                    f"sa_{si}", height=90,
                    placeholder="Write your answer here (2-4 sentences)…",
                    key=f"sa_input_{si}", label_visibility="collapsed"
                )
                if ans:
                    st.session_state.short_answers[si] = ans
            else:
                res = st.session_state.sa_results[si] if si < len(st.session_state.sa_results) else {}
                ans = st.session_state.short_answers.get(si, "")
                score = res.get("score", 0)
                s_col = "#16a34a" if score>=7 else ("#ca8a04" if score>=4 else "#dc2626")
                st.markdown(f'<p style="font-size:.85rem;color:#374151;background:#f9fafb;border-radius:8px;padding:.5rem .75rem">{ans or "(no answer)"}</p>', unsafe_allow_html=True)
                st.markdown(f'<div style="display:flex;align-items:center;gap:.75rem;margin:.4rem 0"><span style="font-weight:700;color:{s_col};font-size:.92rem">{score}/10</span><span style="font-size:.82rem;color:#374151">{res.get("feedback","")}</span></div>', unsafe_allow_html=True)
                if res.get("improvements"):
                    st.markdown(f'<p style="font-size:.78rem;color:#7c3aed;margin:0">💡 Improve: {res["improvements"][0]}</p>', unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)

        if not st.session_state.sa_submitted:
            if st.button("Submit Short Answers →"):
                if api_key:
                    results = []
                    with st.spinner("Scoring short answers with LLM…"):
                        for si, q in enumerate(sas):
                            ans = st.session_state.short_answers.get(si, "")
                            res = score_short_answer(
                                q["question"], q.get("sample_answer",""),
                                q.get("rubric",""), ans, api_key
                            )
                            results.append(res)
                    st.session_state.sa_results   = results
                    st.session_state.sa_submitted = True
                    avg = round(sum(r.get("score",0) for r in results) / len(results), 1)
                    st.success(f"Short answers evaluated — avg score: {avg}/10")
                    st.rerun()
                else:
                    st.warning("An Anthropic API key is required to evaluate short answers.")

    # FINAL ANALYSIS — synthesises ALL pipeline outputs
    all_submitted = (
        bool(st.session_state.task_eval) and
        (not mcqs  or st.session_state.mcq_submitted) and
        (not sas   or st.session_state.sa_submitted)
    )

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="sec-label">Section D · Final Intelligence Report</p>', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">🏁 Complete Candidate Analysis</p>', unsafe_allow_html=True)

    if not st.session_state.task_eval:
        st.markdown("""
        <div style="background:#f9fafb;border:1px dashed #d1d5db;border-radius:12px;padding:1.5rem;text-align:center">
          <p style="color:#9ca3af;font-size:.9rem;margin:0">Complete the assessment above to generate the final analysis.</p>
        </div>""", unsafe_allow_html=True)
    else:
        # Generate/cache final analysis
        if st.session_state.final_analysis is None or st.button("🔄 Refresh Analysis"):
            # Bias check
            dem_fields = {k:v for k,v in {"Name":dem_name,"University":dem_univ,"Grad Year":dem_year}.items() if v.strip()}
            bias_passed = True
            j_s = [s.strip() for s in job_desc.split(",") if s.strip()]
            if dem_fields and (st.session_state.c_skills + st.session_state.bg_skills):
                audit = run_bias_audit(j_s, st.session_state.c_skills + st.session_state.bg_skills, dem_fields)
                bias_passed = audit["all_pass"]
            st.session_state.bias_passed = bias_passed

            fd = compute_final_score(
                st.session_state.match_score,
                st.session_state.task_eval.get("task_score", 0),
                st.session_state.github_data,
                st.session_state.attempts,
                st.session_state.task_eval.get("time_taken", 0),
            )

            with st.spinner("🧠 Synthesising final analysis…"):
                st.session_state.final_analysis = generate_final_analysis(
                    candidate_name=candidate_name,
                    job_title=job_title or (st.session_state.selected_job or {}).get("title",""),
                    match_score=st.session_state.match_score,
                    evidence_chain=st.session_state.evidence_chain,
                    github_data=st.session_state.github_data,
                    task_eval=st.session_state.task_eval,
                    mcq_results=st.session_state.mcq_results,
                    short_answer_results=st.session_state.sa_results,
                    bias_passed=bias_passed,
                    final_score_data=fd,
                    api_key=api_key,
                )

        fa = st.session_state.final_analysis
        if fa:
            fs       = fa["overall_score"]
            hire_rec = fa["hire_recommendation"]
            hire_colors = {
                "STRONG YES": ("#dcfce7","#15803d","✅"),
                "YES":         ("#f0fdf4","#16a34a","✅"),
                "MAYBE":       ("#fef9c3","#ca8a04","⚠️"),
                "NO":          ("#fee2e2","#dc2626","❌"),
            }
            h_bg, h_col, h_icon = hire_colors.get(hire_rec, ("#f9fafb","#374151","•"))
            fs_col = "#4ade80" if fs>75 else ("#fbbf24" if fs>50 else "#f87171")
            conf_cls = fa.get("confidence","MEDIUM").lower()

            # Header row
            st.markdown(f"""
            <div class="final-card">
              <div style="display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:1.5rem;margin-bottom:1.75rem">
                <div>
                  <p style="font-size:.75rem;color:#94a3b8;font-weight:600;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.3rem">Overall Signal Score</p>
                  <div class="final-score" style="color:{fs_col}">{fs}<span style="font-size:1.5rem;color:#475569">/100</span></div>
                  <div style="display:flex;gap:.5rem;margin-top:.6rem;flex-wrap:wrap">
                    <span class="badge {conf_cls}">Confidence: {fa['confidence']}</span>
                    <span class="badge behavior">{fa.get('behavior_profile','')}</span>
                  </div>
                </div>
                <div style="background:{h_bg};border:2px solid {h_col};border-radius:14px;padding:1.1rem 1.75rem;text-align:center;min-width:160px">
                  <div style="font-size:1.5rem">{h_icon}</div>
                  <div style="font-weight:800;font-size:1.1rem;color:{h_col};margin-top:.3rem">HIRE: {hire_rec}</div>
                  <div style="font-size:.72rem;color:{h_col};opacity:.75;margin-top:.15rem">Recommendation</div>
                </div>
              </div>

              <!-- Executive Summary -->
              <div style="background:rgba(255,255,255,.06);border:1px solid rgba(255,255,255,.14);border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1.1rem">
                <p style="font-size:.72rem;font-weight:700;color:#94a3b8;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.5rem">🧠 Executive Summary</p>
                <p style="font-size:.9rem;color:#e2e8f0;line-height:1.75;margin:0">{fa.get('executive_summary','')}</p>
              </div>

              <!-- Signal breakdown -->
              <div style="background:rgba(255,255,255,.04);border:1px solid rgba(255,255,255,.1);border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1.1rem">
                <p style="font-size:.72rem;font-weight:700;color:#94a3b8;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.8rem">📊 Signal Breakdown</p>
                <div style="display:grid;grid-template-columns:repeat(3,1fr);gap:.75rem">""", unsafe_allow_html=True)

            sb = fa.get("signal_breakdown", {})
            breakdown_items = [
                ("🔗 Skill Match",    sb.get("skill_match",0),    "/100"),
                ("💻 Coding Task",    sb.get("coding_task",0),    "/100"),
                ("❓ MCQ Score",      sb.get("mcq",0),            "%"),
                ("✍️ Short Answer",  sb.get("short_answer",0),   "%"),
                ("🐙 GitHub Bonus",   sb.get("github_bonus",0),   " pts"),
                ("🛡️ Bias Check",    "PASS" if sb.get("bias_passed",True) else "FAIL", ""),
            ]
            # Build entire grid as one HTML string — DO NOT loop st.markdown inside open div
            grid_cells = ""
            for label, val, unit in breakdown_items:
                is_str = isinstance(val, str)
                if is_str:
                    col_v = "#4ade80" if val == "PASS" else "#f87171"
                elif val >= 70:
                    col_v = "#4ade80"
                elif val >= 40:
                    col_v = "#fbbf24"
                else:
                    col_v = "#f87171"
                grid_cells += (
                    f'<div style="background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.1);' +
                    f'border-radius:10px;padding:.7rem .75rem;text-align:center;">' +
                    f'<div style="font-size:1.15rem;font-weight:800;color:{col_v};line-height:1">{val}{unit}</div>' +
                    f'<div style="font-size:.72rem;color:#94a3b8;margin-top:.3rem;font-weight:500">{label}</div></div>'
                )
            st.markdown(grid_cells + "</div></div>", unsafe_allow_html=True)

            # Strengths + Gaps side by side (light theme — renders on white bg)
            sc1, sc2 = st.columns(2, gap="medium")

            with sc1:
                strengths_html = """<div style="background:#f0fdf4;border:1px solid #86efac;border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1rem">
  <p style="font-size:.72rem;font-weight:700;color:#16a34a;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.75rem">💪 Top Verified Strengths</p>"""
                for st_ in fa.get("top_strengths", []):
                    src_icon = {"github":"🐙","resume":"📄","task":"💻","dataset":"📦","semantic":"🔗"}.get(st_.get("source",""),"📎")
                    skill_e  = html.escape(st_.get("skill",""))
                    evid_e   = html.escape(st_.get("evidence",""))
                    strengths_html += (
                        f'<div style="display:flex;gap:.6rem;padding:.55rem 0;border-bottom:1px solid #dcfce7">' +
                        f'  <span style="font-size:1rem;flex-shrink:0">{src_icon}</span>' +
                        f'  <div><span style="color:#111827;font-weight:600;font-size:.88rem">{skill_e}</span>' +
                        f'  <p style="font-size:.78rem;color:#6b7280;margin:.1rem 0 0 0">{evid_e}</p></div></div>'
                    )
                strengths_html += "</div>"
                st.markdown(strengths_html, unsafe_allow_html=True)

            with sc2:
                gaps_html = """<div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1rem">
  <p style="font-size:.72rem;font-weight:700;color:#ea580c;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.75rem">🎯 Critical Gaps & Upskill Path</p>"""
                for gap in fa.get("critical_gaps", []):
                    sev      = gap.get("severity","MEDIUM")
                    sev_bg   = "#fee2e2" if sev=="HIGH" else "#fef9c3"
                    sev_col  = "#dc2626" if sev=="HIGH" else "#ca8a04"
                    skill_e  = html.escape(gap.get("skill",""))
                    sugg_e   = html.escape(gap.get("upskill_suggestion",""))
                    gaps_html += (
                        f'<div style="padding:.55rem 0;border-bottom:1px solid #fde8d0">' +
                        f'  <div style="display:flex;align-items:center;gap:.4rem;margin-bottom:.15rem">' +
                        f'    <span style="color:#111827;font-weight:600;font-size:.88rem">{skill_e}</span>' +
                        f'    <span style="background:{sev_bg};color:{sev_col};font-size:.68rem;font-weight:700;padding:.1rem .4rem;border-radius:12px">{sev}</span>' +
                        f'  </div>' +
                        f'  <p style="font-size:.78rem;color:#6b7280;margin:0">📚 {sugg_e}</p></div>'
                    )
                if not fa.get("critical_gaps"):
                    gaps_html += '<p style="color:#16a34a;font-size:.86rem;margin:0">No critical gaps — all required skills covered!</p>'
                gaps_html += "</div>"
                st.markdown(gaps_html, unsafe_allow_html=True)

            # Risk flags
            if fa.get("risk_flags"):
                flags_html = """<div style="background:#fffbeb;border:1px solid #fde68a;border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1rem">
  <p style="font-size:.72rem;font-weight:700;color:#b45309;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.6rem">⚠️ Risk Flags</p>"""
                for flag in fa.get("risk_flags", []):
                    flags_html += f'<p style="font-size:.84rem;color:#374151;margin:.3rem 0">• {html.escape(flag)}</p>'
                flags_html += "</div>"
                st.markdown(flags_html, unsafe_allow_html=True)

            # Next steps
            steps_html = """<div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:1.1rem 1.25rem;margin-bottom:1rem">
  <p style="font-size:.72rem;font-weight:700;color:#1d4ed8;text-transform:uppercase;letter-spacing:.1em;margin-bottom:.6rem">🚀 Recommended Next Steps</p>"""
            for ns in fa.get("next_steps", []):
                steps_html += f'<p style="font-size:.86rem;color:#1e3a5f;margin:.35rem 0;display:flex;gap:.5rem"><span style="color:#4f46e5;font-weight:700">→</span><span>{html.escape(ns)}</span></p>'
            steps_html += "</div>"
            st.markdown(steps_html, unsafe_allow_html=True)

            # Bias footer
            bp = fa["signal_breakdown"].get("bias_passed", True)
            bp_bg  = "#f0fdf4" if bp else "#fff7ed"
            bp_brd = "#86efac" if bp else "#fed7aa"
            bp_col = "#15803d" if bp else "#c2410c"
            bp_txt = "✅ <strong>Bias audit passed.</strong> Score uses only skill artefacts — zero demographic features." if bp else "⚠️ <strong>Score variance detected</strong> when demographic fields removed."
            st.markdown(f'<div style="padding:.85rem 1.1rem;background:{bp_bg};border:1px solid {bp_brd};border-radius:10px"><span style="font-size:.84rem;color:{bp_col}">{bp_txt}</span></div>', unsafe_allow_html=True)

        else:
            st.warning("Please write your code solution and submit to generate the final analysis.")
